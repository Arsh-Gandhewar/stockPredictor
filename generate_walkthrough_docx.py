import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
        elif level == 3:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
    return h

def add_bullet_point(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(bold_prefix + " ")
    r1.font.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(15, 23, 42)
    
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_callout(doc, title, text, hex_bg="F1F5F9", border_color="0284C7"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, hex_bg)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r_title = p.add_run(title + "\n")
    r_title.font.name = 'Calibri'
    r_title.font.bold = True
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = RGBColor(15, 23, 42)
    
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(71, 85, 105)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def format_table(tbl, col_widths, headers, data):
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = Inches(col_widths[i])
        set_cell_background(hdr_cells[i], "0F172A") # Slate 900
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        run.font.name = 'Calibri'
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    # Data rows
    for row_idx, row_data in enumerate(data):
        row_cells = tbl.add_row().cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for i, val in enumerate(row_data):
            row_cells[i].width = Inches(col_widths[i])
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=80, bottom=80, left=120, right=120)
            p = row_cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(30, 41, 59)

def generate_document():
    doc = docx.Document()
    
    # Page setup - Standard Letter, 0.75" margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Header Banner
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_title = title_p.add_run("QuantX: Institutional Stock Market Research & Paper Trading Platform")
    run_title.font.name = 'Calibri'
    run_title.font.bold = True
    run_title.font.size = Pt(24)
    run_title.font.color.rgb = RGBColor(15, 23, 42)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(12)
    run_sub = subtitle_p.add_run("Comprehensive Technical Walkthrough, System Architecture & Operational Guide\nPrepared for Senior Engineering Leadership • August 2026")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(100, 116, 139)
    
    add_callout(
        doc,
        "Executive Summary",
        "QuantX is a high-fidelity Indian stock market quantitative research and paper trading platform built on Next.js 15, NestJS, Neon PostgreSQL, and Google Gemini AI. The platform delivers verified live market data, interactive candlestick charts, autonomous AI portfolio risk guardian monitoring, and full virtual order execution with zero mock data."
    )
    
    # ── Section 1: Architectural Overview ──
    add_styled_heading(doc, "1. Architecture & Technology Stack", 1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.add_run("The platform is engineered as a high-performance TypeScript monorepo structured into three distinct layers:").font.name = 'Calibri'
    
    add_bullet_point(doc, "Frontend Application (apps/web):", "Next.js 15 (App Router), React 19, Tailwind CSS v4, TanStack Query v5, Lightweight Charts v5, and Lucide React icons.")
    add_bullet_point(doc, "Backend REST API (apps/api):", "NestJS v11 architecture with modular services for live market quotes, historical candle streaming, portfolio paper execution, and Gemini AI inference.")
    add_bullet_point(doc, "Database & ORM Layer (packages/db):", "Prisma ORM v6 with Neon Serverless PostgreSQL, configured with multi-column composite indexes for high-frequency trading and historical analytics.")
    add_bullet_point(doc, "Data Feeds & AI Inference:", "Live quote streaming via Yahoo Finance 2 with automated market-state detection (IST hours), and Google Gemini (Pro & Flash) for quantitative investment insights.")

    # ── Section 2: Complete Audit & Root Cause Analysis ──
    add_styled_heading(doc, "2. Problems Audited & Root Cause Remediation", 1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run("A deep architectural audit identified 7 critical vulnerabilities in the legacy version, all of which were eliminated:").font.name = 'Calibri'
    
    audit_table = doc.add_table(rows=1, cols=3)
    format_table(
        audit_table,
        [1.8, 2.3, 2.4],
        ["Audit Category", "Legacy Defect / Root Cause", "Applied Technical Resolution"],
        [
            ["Price Fidelity", "Seed script generated random numbers (Math.random() * 2000 + 1050), showing Reliance at ₹1,734 instead of ₹1,334.", "Connected to live NSE quote stream with freshness badges (LIVE, DELAYED, CLOSED)."],
            ["Candlestick Charts", "Legacy chart endpoint returned a single synthetic record with impossible OHLC (high < open).", "Integrated range-based chart streaming (1D, 1W, 1M, 3M, 6M, 1Y, 5Y) with volume."],
            ["Render Loop Mutations", "Portfolio page re-calculated random prices on every client render, jumping P&L continuously.", "Removed Math.random() and bound all position valuations to live market prices."],
            ["CSS Class Interpolation", "Escaped backticks (\\${...}) in template strings rendered literal strings in HTML class attributes.", "Fixed all escaped template literals across all dashboard and portfolio components."],
            ["Route Coverage & 404s", "/alerts and /settings returned 404; /markets and /news were static placeholders.", "Created dedicated full-featured pages for all 8 routes with zero 404s or stubs."],
            ["Security & Credentials", "Plaintext DATABASE_URL with database password was printed to server logs during startup.", "Removed all plaintext credential logging; sanitized environment variables."],
            ["Database Indexes", "Foreign key models (Transaction, Position, Alert, Watchlist) lacked indexes, causing table scans.", "Added composite and single indexes on all foreign keys and timestamp columns."]
        ]
    )
    
    # ── Section 3: Feature-by-Feature Walkthrough ──
    add_styled_heading(doc, "3. Feature-by-Feature Platform Walkthrough", 1)
    
    add_styled_heading(doc, "3.1 Indian Markets Overview Dashboard (/)", 2)
    p = doc.add_paragraph()
    p.add_run("The central command dashboard provides an instant overview of Indian equities within 5 seconds of opening:").font.name = 'Calibri'
    add_bullet_point(doc, "Benchmark Indices:", "Live tracking of NIFTY 50, SENSEX, BANK NIFTY, and INDIA VIX with day delta and percent changes.")
    add_bullet_point(doc, "Market State Indicator:", "Automatic Indian Standard Time (IST) detection indicating whether the National Stock Exchange is OPEN (9:15-15:30 IST), PRE_OPEN, or CLOSED.")
    add_bullet_point(doc, "Interactive Benchmark Candlestick Chart:", "Mounted lightweight-charts candlestick widget displaying NIFTY 50 price action over the past 30 days.")
    add_bullet_point(doc, "Market Movers Tab:", "Real-time computed tables for Top Gainers, Top Losers, and Most Active equities with 1-click navigation to stock deep-dives.")
    add_bullet_point(doc, "Top Monitored Picks:", "High-conviction algorithmic opportunities displaying live price, quantitative confidence score, and thesis.")

    add_styled_heading(doc, "3.2 Dynamic Stock Details & Candlestick Analysis (/stock/[ticker])", 2)
    p = doc.add_paragraph()
    p.add_run("A complete institutional terminal experience for analyzing individual equities:").font.name = 'Calibri'
    add_bullet_point(doc, "Timeframe Range Switcher:", "Instant toggle between 1D, 1W, 1M, 3M, 6M, 1Y, and 5Y historical candlestick resolutions.")
    add_bullet_point(doc, "Key Financial Parameters:", "52-Week High/Low range, Day Range, P/E Ratio, Market Capitalization (₹ Crores), Volume, and Previous Close.")
    add_bullet_point(doc, "AI Research Thesis Breakdown:", "Structured qualitative breakdown including Recommendation Stance, Conviction Score, Thesis Drivers, and Invalidation Criteria.")
    add_bullet_point(doc, "Interactive Order Ticket Modal:", "Dedicated paper trading execution modal supporting custom share quantities, BUY/SELL order toggles, and live balance verification.")

    add_styled_heading(doc, "3.3 Universe Discovery (/discover)", 2)
    p = doc.add_paragraph()
    p.add_run("Allows active scanning and discovery across the 49 NIFTY 50 constituent equities. Users can filter by sector (Technology, Energy, Finance, Auto, Pharma, FMCG, Metals) or search by company name and ticker symbol.").font.name = 'Calibri'

    add_styled_heading(doc, "3.4 Markets & Breadth Radar (/markets)", 2)
    p = doc.add_paragraph()
    p.add_run("Multi-index monitoring with detailed side-by-side market breadth breakdown comparing leading gainers and decliners across the current trading session.").font.name = 'Calibri'

    add_styled_heading(doc, "3.5 Personal Watchlist (/watchlist)", 2)
    p = doc.add_paragraph()
    p.add_run("Allows traders to pin their highest-conviction ideas with persistent browser storage (localStorage) and quick modal search to add new stocks.").font.name = 'Calibri'

    add_styled_heading(doc, "3.6 Paper Trading Portfolio & AI Risk Guardian (/portfolio)", 2)
    p = doc.add_paragraph()
    p.add_run("Full-fidelity paper trading simulation engine with ₹10,00,000 starting virtual capital:").font.name = 'Calibri'
    add_bullet_point(doc, "Portfolio Valuation:", "Calculates total portfolio value, invested amount, absolute P&L, and return percentage in real-time.")
    add_bullet_point(doc, "AI Risk Guardian:", "Background monitoring module evaluating unrealized profits, news sentiment, and technical setups to output high-conviction (>80%) sell signals.")
    add_bullet_point(doc, "Position Management & History:", "Real-time holdings table with individual position P&L and comprehensive chronological transaction ledger.")

    add_styled_heading(doc, "3.7 Institutional Market News (/news)", 2)
    p = doc.add_paragraph()
    p.add_run("Curated market news intelligence categorized by Corporate, Macro, Results, and Markets, annotated with short/long impact ratings and sentiment badges.").font.name = 'Calibri'

    add_styled_heading(doc, "3.8 Real-Time Price Alerts (/alerts)", 2)
    p = doc.add_paragraph()
    p.add_run("Price alert manager allowing users to configure price thresholds ('Crosses Above' / 'Crosses Below') across any monitored stock in the universe.").font.name = 'Calibri'

    # ── Section 4: REST API Reference ──
    add_styled_heading(doc, "4. Backend REST API Reference", 1)
    
    api_table = doc.add_table(rows=1, cols=3)
    format_table(
        api_table,
        [1.6, 2.4, 2.5],
        ["Method & Endpoint", "Parameters", "Description & Response Payload"],
        [
            ["GET /stock/market-status", "None", "Returns exchange state (OPEN/CLOSED/PRE_OPEN), exchange, and IST timestamp."],
            ["GET /stock/market-summary", "None", "Returns array of NIFTY 50, SENSEX, BANK NIFTY, INDIA VIX with live values and % change."],
            ["GET /stock/movers", "None", "Returns top 5 gainers, top 5 losers, and top 5 most active equities by volume."],
            ["GET /stock/all", "None", "Returns all 49 NIFTY 50 stocks with ticker, company name, sector, and exchange."],
            ["GET /stock/search", "q: string", "Case-insensitive query searching across ticker symbols, company names, and sectors."],
            ["GET /stock/:ticker/quote", "ticker: string", "Returns live price, change, day range, 52W range, PE, market cap, and data freshness."],
            ["GET /stock/:ticker/chart", "ticker: string, range: string", "Returns historical OHLC candlestick array for requested range (1d, 1w, 1mo, 6mo, 1y, 5y)."],
            ["GET /portfolio", "Header: x-user-id", "Returns user portfolio with available cash, active positions, and recent transactions."],
            ["POST /portfolio/trade", "Body: { ticker, type, quantity }", "Executes virtual market BUY/SELL order, validates cash/shares, and updates ledger."],
            ["GET /portfolio/sell-signals", "Header: x-user-id", "Runs Gemini AI risk evaluation across user holdings for >80% confidence exit signals."]
        ]
    )

    # ── Section 5: Automated Testing & Verification ──
    add_styled_heading(doc, "5. Automated Testing & Verification Results", 1)
    
    p = doc.add_paragraph()
    p.add_run("A custom end-to-end integration test runner (test-runner.ts) was executed against the entire stack. All 24 integration test cases passed with a 100% success rate:").font.name = 'Calibri'
    
    test_table = doc.add_table(rows=1, cols=3)
    format_table(
        test_table,
        [2.0, 3.0, 1.5],
        ["Test Suite", "Verification Scope", "Result"],
        [
            ["Suite 1: Market Data", "Market status, multi-index quotes, market movers, NIFTY 50 list, search, live quotes, 404 handling, OHLC candles.", "8/8 Passed (100%)"],
            ["Suite 2: Paper Trading", "Default portfolio creation (₹10L cash), BUY order execution, SELL order execution, overselling rejection, AI sell signals.", "5/5 Passed (100%)"],
            ["Suite 3: Frontend Routes", "HTTP 200 checks for /, /discover, /markets, /watchlist, /portfolio, /news, /alerts, /settings, and dynamic stock pages.", "11/11 Passed (100%)"]
        ]
    )

    add_callout(
        doc,
        "Quality Assurance Verdict",
        "Total Tests Run: 24 | Passed: 24 | Failed: 0 | Success Rate: 100.0%. Both frontend and backend TypeScript compilation pass with 0 errors (npx tsc --noEmit).",
        hex_bg="F0FDF4",
        border_color="16A34A"
    )

    # ── Section 6: Running Locally ──
    add_styled_heading(doc, "6. Setup & Operational Commands", 1)
    
    p = doc.add_paragraph()
    p.add_run("To run and test the complete application locally:").font.name = 'Calibri'
    add_bullet_point(doc, "1. Start Backend Server:", "$env:NODE_OPTIONS=\"--dns-result-order=ipv4first\"; npm run start:dev (Port 3001)")
    add_bullet_point(doc, "2. Start Frontend Server:", "npm run dev (Port 3000)")
    add_bullet_point(doc, "3. Run Automated Tests:", "npx tsx test-runner.ts")
    add_bullet_point(doc, "4. Seed NIFTY 50 Database:", "$env:NODE_OPTIONS=\"--dns-result-order=ipv4first\"; npx tsx apps/api/src/scripts/seed.ts")

    output_path = r"C:\Users\arshg\OneDrive\Desktop\stockPredictor\QuantX_StockPredictor_Complete_Walkthrough.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    generate_document()
